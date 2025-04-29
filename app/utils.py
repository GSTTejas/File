import os
from docx import Document

SPLIT_DIR = "media/split_docs"
MERGED_DIR = "media/merged_docs"

def split_word_document(input_docx, parts=10):
    """Splits a Word document into equal parts and saves them."""
    if not os.path.exists(input_docx):
        print(f"❌ File not found: {input_docx}")
        return False

    os.makedirs(SPLIT_DIR, exist_ok=True)

    try:
        doc = Document(input_docx)
        print(f"✅ Loaded DOCX file: {input_docx}")
    except Exception as e:
        print(f"❌ Error opening DOCX: {e}")
        return False

    paragraphs = doc.paragraphs
    num_paragraphs = len(paragraphs)

    if num_paragraphs == 0:
        print("❌ No paragraphs found in the document!")
        return False

    # Avoid empty parts if paragraphs < parts
    if num_paragraphs < parts:
        parts = num_paragraphs

    paragraphs_per_part = num_paragraphs // parts
    remaining_paragraphs = num_paragraphs % parts

    start_paragraph = 0

    for part in range(parts):
        new_doc = Document()
        end_paragraph = start_paragraph + paragraphs_per_part

        if part < remaining_paragraphs:
            end_paragraph += 1

        # Copy paragraphs into new document
        for i in range(start_paragraph, end_paragraph):
            new_doc.add_paragraph(paragraphs[i].text)

        output_filename = os.path.join(SPLIT_DIR, f"split_part_{part + 1}.docx")
        new_doc.save(output_filename)
        print(f"✅ Saved split file: {output_filename}")

        start_paragraph = end_paragraph

    print("✅ Splitting process completed successfully!")
    return True

def merge_documents(file_paths, output_file):
    """Merges multiple DOCX files into a single document."""
    if not file_paths:
        print("❌ No files provided for merging!")
        return False

    os.makedirs(MERGED_DIR, exist_ok=True)
    
    merged_doc = Document()
    for file in file_paths:
        try:
            temp_doc = Document(file)
            for para in temp_doc.paragraphs:
                merged_doc.add_paragraph(para.text)
            if file != file_paths[-1]:
                merged_doc.add_page_break()
        except Exception as e:
            print(f"❌ Error merging {file}: {e}")
            return False

    try:
        merged_doc.save(output_file)
        print(f"✅ Merged document saved at: {output_file}")
        return True
    except Exception as e:
        print(f"❌ Error saving merged document: {e}")
        return False

# pdf

import os
from pypdf import PdfReader, PdfWriter
from django.conf import settings

UPLOAD_DIR = os.path.join(settings.MEDIA_ROOT, "uploads")
SPLIT_DIR = os.path.join(settings.MEDIA_ROOT, "split_pdfs")
MERGED_DIR = os.path.join(settings.MEDIA_ROOT, "merged_pdfs")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(SPLIT_DIR, exist_ok=True)
os.makedirs(MERGED_DIR, exist_ok=True)

def save_uploaded_pdf(pdf_file):
    """Saves the uploaded PDF to the media directory."""
    file_path = os.path.join(UPLOAD_DIR, pdf_file.name)
    with open(file_path, "wb") as f:
        for chunk in pdf_file.chunks():
            f.write(chunk)
    return file_path

def split_pdf_in_parts(input_pdf, parts=10):
    """Splits a PDF into equal parts and saves them."""
    reader = PdfReader(input_pdf)
    num_pages = len(reader.pages)
    parts = min(parts, num_pages)  # Avoid empty files
    pages_per_part = num_pages // parts
    remaining_pages = num_pages % parts
    start_page = 0

    for part in range(parts):
        writer = PdfWriter()
        end_page = start_page + pages_per_part + (1 if part < remaining_pages else 0)

        for page_num in range(start_page, end_page):
            writer.add_page(reader.pages[page_num])

        output_filename = os.path.join(SPLIT_DIR, f"split_part_{part + 1}.pdf")
        with open(output_filename, 'wb') as output_pdf:
            writer.write(output_pdf)

        start_page = end_page

def merge_pdfs():
    """Merges all split PDFs into a single file."""
    merged_file_path = os.path.join(MERGED_DIR, "merged_pdf.pdf")
    writer = PdfWriter()

    files = sorted([os.path.join(SPLIT_DIR, f) for f in os.listdir(SPLIT_DIR) if f.endswith(".pdf")])
    if not files:
        return None

    for file in files:
        reader = PdfReader(file)
        for page in reader.pages:
            writer.add_page(page)

    with open(merged_file_path, "wb") as merged_pdf:
        writer.write(merged_pdf)

    return merged_file_path

def verify_pdf_pin(pdf, entered_pin):
    """Verifies if the entered PIN matches the stored PIN."""
    return pdf.pin == entered_pin

from .models import UploadedPDF

def verify_pdf_pin(pdf, entered_pin):
    """Verifies the PIN before allowing PDF download."""
    return pdf.security_pin == entered_pin  # ✅ Compare stored PIN


# videos

import os
import subprocess
import logging
from django.conf import settings

logger = logging.getLogger(__name__)

def get_video_duration(input_video_path):
    """
    Gets the duration of the video in seconds using FFprobe.
    """
    cmd = [
        settings.FFPROBE_PATH,
        "-v", "error",
        "-select_streams", "v:0",
        "-show_entries", "format=duration",
        "-of", "default=noprint_wrappers=1:nokey=1",
        input_video_path
    ]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        return float(result.stdout.strip())
    except subprocess.CalledProcessError as e:
        logger.error(f"FFprobe failed: {e}")
        return None
    except ValueError:
        logger.error("Could not determine video duration.")
        return None

def split_video(input_video_path, output_folder):
    """
    Splits a video into 10 equal parts efficiently using FFmpeg's segment filter.
    Returns a list of split file paths.
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    total_seconds = get_video_duration(input_video_path)
    if total_seconds is None:
        logger.error("Failed to get video duration. Aborting split process.")
        return None

    part_duration = total_seconds / 10  # Each part duration
    output_template = os.path.join(output_folder, "part_%02d.mp4")

    cmd_split = [
    settings.FFMPEG_PATH,
    "-i", input_video_path,  
    "-c", "copy",
    "-map", "0:v:0",  # Select only the first video stream
    "-map", "0:a:0?",  # Select the first audio stream (if available)
    "-f", "segment",
    "-segment_time", str(part_duration),
    "-reset_timestamps", "1",
    output_template
]


    try:
        subprocess.run(cmd_split, check=True)
        split_files = [os.path.join(output_folder, f) for f in sorted(os.listdir(output_folder)) if f.endswith(".mp4")]
        return split_files if split_files else None

    except subprocess.CalledProcessError as e:
        logger.error(f"Error splitting video: {e}")
        return None

def merge_videos(split_files, output_path):
    """
    Merges a list of split video files into a single video using FFmpeg.
    """
    if not split_files:
        logger.error("No split files provided for merging.")
        return None

    output_dir = os.path.dirname(output_path)
    os.makedirs(output_dir, exist_ok=True)  

    temp_txt = os.path.join(output_dir, "merge_list.txt")

    try:
        # Create a list of files for FFmpeg
        with open(temp_txt, "w", encoding="utf-8") as f:
            for file in sorted(split_files):  
                f.write(f"file '{os.path.abspath(file)}'\n")

        # Run FFmpeg merge command
        cmd_merge = [
            settings.FFMPEG_PATH,
            "-f", "concat",
            "-safe", "0",
            "-i", temp_txt,
            "-c", "copy",  # Fast merging without re-encoding
            output_path
        ]
        subprocess.run(cmd_merge, check=True)

        os.remove(temp_txt)  # Clean up temp file

        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            logger.info(f"Merged video successfully created: {output_path}")
            return output_path
        else:
            logger.error("Merging failed: Output file is empty.")
            return None

    except subprocess.CalledProcessError as e:
        logger.error(f"Error merging videos: {e}")
        return None
 
